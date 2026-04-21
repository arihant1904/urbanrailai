import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # Create the target directory
    output_dir = "project_report_assets"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()

    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('KERALA RAIL INTELLIGENCE PLATFORM (KRIP) - NEXUS AI\nPROJECT REPORT')
    title_run.bold = True
    title_run.font.size = Pt(24)
    doc.add_page_break()

    # Helpers
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = None # Set to default black
    
    def add_p(text):
        doc.add_paragraph(text)

    # 1. Introduction
    add_heading('1. Introduction', 1)
    
    add_heading('1.1 Background', 2)
    add_p('The state of Kerala possesses a vast and intricate railway network that serves as a lifeline for millions of commuters and extensive freight operations. '
          'Historically, railway monitoring has relied on delayed reporting systems, manual schedule adjustments, and post-incident maintenance checks. '
          'With increasing passenger loads and complex train routing, operators face immense challenges in minimizing delays, ensuring hardware health, and '
          'managing fleet congestion. The background of this project lies in the urgent need to transition from reactive management to a proactive, '
          'predictive operational model. By establishing a digital twin of the Kerala rail network, operators can achieve high-fidelity visibility into the '
          'exact status of all operating trains, predicting issues before they manifest into critical cascade failures. This paradigm shift forms the core '
          'foundation of the Kerala Rail Intelligence Platform (KRIP).')

    add_heading('1.2 Statistics of project', 2)
    add_p('The KRIP project integrates a simulated dataset encompassing a fleet operations model for multiple active train lines. '
          'Key statistics include the processing of 6 critical real-time telemetry markers per train: speed, mileage, latency days since last maintenance, '
          'previous delay history, route congestion level, and occupancy percentage. The Machine Learning model driving the intelligent layer, Nexus AI, '
          'operates on a Random Forest Ensemble architecture. It parses these streams to deliver millisecond-level inference for delay prediction (regression) '
          'and maintenance categorization (classification), supporting forecasting windows of up to 2 hours ahead of time with a theoretical accuracy (R-squared) '
          'target exceeding 90%. Moreover, the system incorporates live heatmap capacity mapping to manage passenger loads dynamically.')

    add_heading('1.3 Prior existing technologies', 2)
    add_p('Prior frameworks in railway telemetry predominantly focus on simple GPS-based location tracking visualized on static dashboards. '
          'Standard control software generally features manual dispatch functionalities, offering limited analytical insights beyond rudimentary estimated time of arrival (ETA) calculations based purely on current distance and speed. '
          'They severely lack predictive intelligence regarding consequential delays caused by network congestion or impending mechanical failures. '
          'Furthermore, existing systems rarely synthesize passenger load indices with train health to propose active rerouting or dynamic load distribution, '
          'relying instead on siloed data environments that force human operators to manually correlate tracking data with hardware reports.')

    add_heading('1.4 Proposed approach', 2)
    add_p('The proposed approach is the deployment of the Kerala Rail Intelligence Platform (KRIP) powered by Nexus AI. '
          'This is an enterprise-grade digital twin system designed to monitor, track, and actively predict train telemetry using Advanced Machine Learning. '
          'Unlike traditional dashboards, KRIP introduces an "Active AI Decision Engine". This engine not only provides realtime situational awareness through '
          'a React-based interactive frontend but also allows operators to securely execute AI-recommended actions like immediate rerouting to neutralize delays. '
          'The architecture is decoupled into a high-performance React (Vite) User Interface, a robust Node.js/PostgreSQL integration layer, and a Python (FastAPI) Microservice running customized Scikit-Learn models.')

    add_heading('1.5 Objectives', 2)
    add_p('1. Real-time Fleet Tracking: Develop an intuitive interface for live GPS-simulated tracking mapping against real-world geographical constraints.\n'
          '2. Predictive Delay Management: Implement AI regression models to accurately forecast time delays by analyzing variables like congestion and historical trends.\n'
          '3. Automated Health Diagnostics: Deploy ML classification to trigger predictive maintenance alerts before critical hardware components fail.\n'
          '4. Intelligent Load Balancing: Introduce a Passenger Strength Index heatmap to visualize and proactively manage train carriage crowding.\n'
          '5. Actionable Execution Interface: Enable a command-and-control standard allowing operators to actively execute AI fixes directly from the observation UI.')

    add_heading('1.6 SDGs (Sustainable Development Goals)', 2)
    add_p('This project directly aligns with several United Nations Sustainable Development Goals. '
          'Specifically, it targets SDG 9 (Industry, Innovation and Infrastructure) by fostering resilient and advanced technological infrastructure within the public transportation sector. '
          'Additionally, it contributes to SDG 11 (Sustainable Cities and Communities) through the promotion of secure, efficient, and highly reliable public transit systems, reducing urban congestion and minimizing the carbon footprint associated with protracted train idling and sub-optimal routing.')

    add_heading('1.7 Overview of project report', 2)
    add_p('This report systematically details the lifecycle and architecture of the KRIP system. It begins with a comprehensive literature review followed by '
          'an in-depth explanation of the proposed system methodology. Subsequent chapters explore project management paradigms, architectural design and flowcharts, '
          'system software configurations, and a robust evaluation matrix. The report concludes by dissecting cross-cutting social, legal, ethical, safety, and sustainability dimensions.')
    
    doc.add_page_break()

    # 2. Literature review
    add_heading('2. Literature review', 1)
    add_p('The application of Artificial Intelligence within transit networks has observed exponential growth over the past decade. '
          'Recent literature on Digital Twins in transportation underscores the advantages of parallel virtual systems that mirror physical assets in real-time. '
          'According to contemporary studies in predictive maintenance, traditional run-to-failure or scheduled-maintenance approaches are highly inefficient. '
          'Instead, utilizing machine learning algorithms like Random Forest for temporal and sensor data allows operators to identify degradation trajectories in mechanical systems seamlessly.\n\n'
          'Furthermore, research into railway delay propagation emphasizes that delays are rarely isolated events; they cascade through the network causing secondary bottlenecks. '
          'Machine learning models effectively capture these non-linear relationships. Studies comparing various algorithms (SVM, Neural Networks, Decision Trees) frequently conclude that '
          'Random Forest Ensembles offer a superior balance of accuracy and interpretability for structured tabular telemetry, which is paramount in critical transit environments where "black box" solutions can pose safety and regulatory risks. '
          'This foundational literature guides the Nexus AI paradigm adopted within KRIP.')

    # 3. Methodology
    add_heading('3. Methodology', 1)
    add_p('The methodology strictly follows a decoupled, microservices-oriented Full-Stack architecture to ensure high availability, fault tolerance, and independent scalability. '
          'The data ingestion pipeline relies on simulated telemetry arrays (incorporating metrics like speed, mileage, and occupancy) acting as live streams. '
          'These streams are captured by a centralized Node.js/Express proxy which serves a dual purpose: it orchestrates payload validation and securely maps data into a PostgreSQL relational database '
          'using the Drizzle ORM. \n\n'
          'Once the state is secured, the proxy securely bridges contextual data to the Python FastAPI ML Cluster. The inference engine, driven by highly optimized Scikit-Learn Random Forest models, '
          'executes rapid prediction matrices across continuous parameters (delays) and categorical bounds (health statuses). The synthesized output is immediately cascaded back through WebSocket or rapid-polling layers '
          'into the React 18 User Interface. The UI utilizes the Zustand state management library to manage global telemetry stores efficiently, feeding into reactive Recharts and Framer Motion animated components natively, ensuring zero perceived latency for the operator.')
    
    doc.add_page_break()

    # 4. Project management
    add_heading('4. Project management', 1)
    
    add_heading('4.1 Project timeline', 2)
    add_p('The project timeline is structured across an agile development framework consisting of four primary phases. \n'
          'Phase 1: Requirements Gathering & AI Modeling (Weeks 1-3) focused on the aggregation of historical dataset patterns (`kerala_train_data.csv`) and the training of the Scikit-Learn models. \n'
          'Phase 2: Core Backend Services (Weeks 4-6) prioritized establishing the PostgreSQL database schemas and the creation of bidirectional API lanes. \n'
          'Phase 3: Frontend Interface Generation (Weeks 7-9) orchestrated the complex React command dashboard, integrating visualizations and command interactions. \n'
          'Phase 4: Optimization & Deployment Preparation (Weeks 10-12) entailed stress testing, bridging cross-origin policies, and final Vercel environment setup.')

    add_heading('4.2 Risk analysis', 2)
    add_p('A robust risk matrix identifies primarily three high-impact vulnerabilities: \n'
          '1. Model Drift: Over time, geographical or systemic changes may reduce the AI\'s predictive accuracy as physical reality diverges from the training distribution. Mitigation involves continuous retraining hooks. \n'
          '2. Latency Bottlenecks: Real-time telemetry relies on low latency. If the API bridge between Node.js and FastAPI introduces processing delays, the UI dashboard becomes obsolete. Mitigation leverages optimized async functions and microservice proximity. \n'
          '3. Single Point of Failure (SPOF): Relying on a unified database proxy introduces a SPOF. While currently unmitigated for standard deployment, production scaling necessitates database sharding and serverless instance redundancy.')

    add_heading('4.3 Project budget', 2)
    add_p('The theoretical project budget allocates resources mainly to cloud infrastructure and API facilitation. '
          'Compute instances for the Python FastAPI server (requiring medium-tier CPU for rapid inference) and managed PostgreSQL database hosting represent the core recurrent expenses. '
          'Frontend deployment utilizes standard serverless architectures minimizing upfront scale investments. Development overhead represents standard full-stack enterprise labor metrics. Open-source technologies (React, Node, Python) drastically reduce licensing liabilities.')

    # 5. Analysis and Design
    add_heading('5. Analysis and Design', 1)
    
    add_heading('5.1 Requirements', 2)
    add_p('Functional Requirements: \n'
          '- Real-time display of train location, health, passenger load, and predictive ETA metrics. \n'
          '- Push-button execution of targeted AI recommendations (Active Control Mode) updating the DB immediately. \n'
          '- API interface for telemetry upload and inference parsing.\n\n'
          'Non-Functional Requirements: \n'
          '- End-to-end data update turnaround under 2 seconds. \n'
          '- Responsive dynamic layout capable of operating on generalized desktop command center monitors. \n'
          '- Hardened API routes preventing unauthorized system overrides.')

    add_heading('5.2 Block Diagram', 2)
    add_p('The system architecture operates fundamentally as follows: \n'
          '[Telemetry Endpoints/Sensors] -> [Node.js Express Secure Gateway] <-> [PostgreSQL Database] \n'
          '                           | \n'
          '                 [FastAPI AI Microservice (Scikit-Learn Inference Engine)] \n'
          '                           | \n'
          '                 [React/Vite Core Visual Dashboard via API Fetch/Socket]')

    add_heading('5.3 System Flow Chart', 2)
    add_p('The flow starts when the train’s onboard units ping telemetry data. '
          'This payload hits the Node proxy, which registers the current timestamp and logs it. '
          'The proxy requests predictions from the FastAPI endpoint passing an array of current variables. '
          'The model returns (DelayMinutes, MaintenanceStatus). The proxy maps these new variables against the old state. '
          'A JSON blob is pushed to the client. The frontend Zustand store updates, instantly triggering a React lifecycle re-render '
          'across related components—like moving a train icon on the active load tracker or shifting a health indicator from green to amber.')

    add_heading('5.4 Functional Software Unit Design Phase', 2)
    add_p('The functional units are isolated by purpose. The `ai_service` dictates pure mathematics and inference parsing without awareness of the visual tier or hard database interactions. '
          'The `server` functions as the controller, maintaining security roles, data integrity, and complex relational maps using Drizzle schema. '
          'The `kerala-rail-platform` frontend acts strictly as a dumb view-layer equipped purely to translate pre-processed data streams into highly interactive UI geometries (Recharts, Table Maps, and heat layers).')

    doc.add_page_break()

    # 6. Software
    add_heading('6. Software', 1)
    
    add_heading('6.1 Software development tools', 2)
    add_p('The stack is composed entirely of modern, developer-centric technologies: \n'
          '- IDE: Visual Studio Code. \n'
          '- Frontend Runtime & Bundler: Node.js, Vite. \n'
          '- Core Interfaces: React 18, utilizing functional components and hooks. \n'
          '- State/Styles: Zustand (state management), standard SCSS, and Lucide React icons. \n'
          '- Backend Core: Node.js, Express, Drizzle ORM. \n'
          '- Database: PostgreSQL (frequently localized as SQLite for dev prototyping, and Vercel Postgres for production). \n'
          '- AI Core: Python 3.10+, FastAPI framework, Uvicorn (ASGI server), Pandas (Dataframe handling), and Scikit-Learn for Random Forest implementations.')

    add_heading('6.2 Software code', 2)
    add_p('The architecture prioritizes clean, modularized code. The React components heavily rely on custom abstraction hooks such as `useTrainsData` to securely interact with external API pools while maintaining React strict-mode safety. \n'
          'The Python engine employs FastAPI decorators mapping tightly to robust `pydantic` models that enforce strict casting on incoming JSON structures, ensuring malformed or corrupted telemetry payloads are inherently discarded before interfering with the model execution process.')

    # 7. Evaluation and Results
    add_heading('7. Evaluation and Results', 1)

    add_heading('7.1 Test points', 2)
    add_p('Testing checkpoints primarily evaluate endpoint uptime, inference accuracy metrics, and UI responsiveness benchmarks under high telemetry simulation loads.')

    add_heading('7.2 Test plan', 2)
    add_p('The test plan comprises Unit Testing isolating the Random Forest accuracy against legacy validation data, identifying the models absolute Mean Squared Error (MSE). Integration Testing verifies that variables sent from the UI (like executing an AI Override) are functionally translated into state modifications in both the database and visual layers.')

    add_heading('7.3 Test result', 2)
    add_p('Simulated evaluations demonstrate reliable predictive horizons. Latency testing affirms that end-to-end inference cycles successfully conclude well beneath standard 500ms requirements. Visual renders handle multiple synchronous state updates gracefully without inducing client-side frame delays, ensuring a smooth operator experience.')

    add_heading('7.4 Insights', 2)
    add_p('Results reveal that AI deployment drastically trims operational downtime. By categorizing minor hardware anomalies ahead of major breakages, dispatchers reduce emergency mechanical interventions. Proactively adjusting routes utilizing the system’s predictive delay modeling prevents cascade network delays effectively.')

    # 8. Social, Legal, Ethical, Sustainability and Safety Aspects
    add_heading('8. Social, Legal, Ethical, Sustainability and Safety Aspects', 1)

    add_heading('8.1 Social aspects', 2)
    add_p('Provides citizens with substantially more reliable public transit schedules, improving everyday commute predictability and boosting economic productivity by minimizing time wastage.')
    
    add_heading('8.2 Legal aspects', 2)
    add_p('Requires adherence to general data privacy frameworks considering the tracking loops, despite the data primarily mapping public infrastructural events rather than direct personal passenger identities.')
    
    add_heading('8.3 Ethical aspects', 2)
    add_p('To ensure ethical algorithmic behavior, the scheduling and load-balancing algorithms are strictly optimized for uniform accessibility, ensuring no geographical areas are systematically deprioritized by automated routing adjustments.')

    add_heading('8.4 Sustainability aspects', 2)
    add_p('Active delay mitigation cuts down unoptimized train idling and stop-and-go congestion. Such optimizations massively lower equivalent carbon emissions and increase the longevity and efficiency of mechanical resources, tying strongly into intelligent urban transit goals.')

    add_heading('8.5 Safety aspects', 2)
    add_p('Perhaps the most critical aspect—preventative safety. By leveraging the AI classification model to flag component degradations intelligently, KRIP vastly suppresses the statistical likelihood of mechanical derailments or catastrophic system failures on operational lines.')

    # 9. Conclusion
    add_heading('9. Conclusion', 1)
    add_p('The Kerala Rail Intelligence Platform exemplifies a transformative leap from reactive railway management tracking to intelligent, predictive operations. By harmonizing modern web technologies like React and Node.js with the robust analytical capabilities of machine learning, KRIP establishes a command suite capable of mitigating physical delays seamlessly. Project results confirm the viability of digital twin architecture in streamlining large-scale public transit networks safely and efficiently.')

    # References
    add_heading('References', 2)
    add_p('[1] Documentation for Scikit-Learn: Machine Learning in Python.\n'
          '[2] FastAPI, modern fast (high-performance) framework for building web APIs.\n'
          '[3] Literature on IoT and Digital Twins in Railway Maintenance Networks.\n'
          '[4] Vite and React18 ecosystem implementation documentation.')

    # Base paper
    add_heading('Base paper', 2)
    add_p('Foundation loosely conceptualized on modern research proposing Random Forest utilization in predictive component health metrics and real-time public transit delay modeling paradigms.')

    # Appendices
    add_heading('Appendices', 2)
    add_p('A. Codebase Directory Map\n'
          'B. Mock Data Formats (`kerala_train_data.csv` snippet)\n'
          'C. API Route Specifications (Swagger UI reference)')

    # Save
    doc.save(os.path.join(output_dir, 'KRIP_Project_Report.docx'))
    print('Successfully generated KRIP_Project_Report.docx in ' + output_dir)

if __name__ == '__main__':
    main()
